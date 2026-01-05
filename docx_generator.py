import os
import json
from datetime import datetime
from docxtpl import DocxTemplate
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_memorial_docx(projeto):
    """Gera um documento DOCX do memorial descritivo fotovoltaico"""
    
    # Caminho do template
    template_path = os.path.join(os.path.dirname(__file__), 'templates', 'modelo_memorial_v2.docx')
    
    # Se o template não existir, criar um documento básico
    if not os.path.exists(template_path):
        return generate_memorial_from_scratch(projeto)
    
    # Carregar template
    doc = DocxTemplate(template_path)
    
    # Preparar contexto para o template
    context = prepare_context(projeto)
    
    # Renderizar template
    doc.render(context)
    
    # Salvar documento
    output_folder = os.path.join(os.path.dirname(__file__), 'uploads')
    os.makedirs(output_folder, exist_ok=True)
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    nome_arquivo = f"Memorial_{projeto['nome_cliente']}_{timestamp}.docx"
    output_path = os.path.join(output_folder, nome_arquivo)
    
    doc.save(output_path)
    
    return output_path

def prepare_context(projeto):
    """Prepara o contexto para renderizar o template Jinja"""
    
    # Parsear equipamentos se forem strings JSON
    modulos_novos = projeto.get('modulos_novos', [])
    if isinstance(modulos_novos, str):
        modulos_novos = json.loads(modulos_novos) if modulos_novos else []
    
    inversores_novos = projeto.get('inversores_novos', [])
    if isinstance(inversores_novos, str):
        inversores_novos = json.loads(inversores_novos) if inversores_novos else []
    
    # Calcular Pg para Art. 73-A
    pg = 0
    if projeto.get('tipo_projeto') == 'Art. 73-A':
        media_consumo = float(projeto.get('media_consumo', 0))
        fator_carga = float(projeto.get('fator_carga', 0))
        fator_ajuste = float(projeto.get('fator_ajuste', 0))
        
        if media_consumo > 0 and fator_carga > 0:
            pg = (media_consumo * fator_carga * fator_ajuste) / 1000
    
    context = {
        # Dados do cliente
        'nome_cliente': projeto.get('nome_cliente', ''),
        'cpf_cnpj': projeto.get('cpf_cnpj', ''),
        'uc': projeto.get('uc', ''),
        'endereco': projeto.get('endereco', ''),
        'cidade': projeto.get('cidade', ''),
        'uf': projeto.get('uf', ''),
        'cep': projeto.get('cep', ''),
        'concessionaria': projeto.get('concessionaria', ''),
        'data_projeto': projeto.get('data_projeto', ''),
        
        # Tipo de projeto
        'tipo_projeto': projeto.get('tipo_projeto', ''),
        'eh_instalacao_nova': projeto.get('tipo_projeto') == 'Instalação Nova',
        'eh_ampliacao': projeto.get('tipo_projeto') == 'Ampliação',
        'eh_grid_zero': projeto.get('tipo_projeto') == 'Grid Zero',
        'eh_art73a': projeto.get('tipo_projeto') == 'Art. 73-A',
        
        # Ampliação
        'modulos_existentes': projeto.get('modulos_existentes', 0),
        'inversores_existentes': projeto.get('inversores_existentes', ''),
        
        # Grid Zero
        'controlador': projeto.get('controlador', ''),
        'transdutor_tc': projeto.get('transdutor_tc', ''),
        'chave_seccionadora': projeto.get('chave_seccionadora', ''),
        
        # Art. 73-A
        'media_consumo': projeto.get('media_consumo', 0),
        'fator_carga': projeto.get('fator_carga', 0),
        'fator_ajuste': projeto.get('fator_ajuste', 0),
        'pg': round(pg, 2),
        
        # Novo Sistema
        'modulos_novos': modulos_novos,
        'inversores_novos': inversores_novos,
        'total_modulos': len(modulos_novos),
        'total_inversores': len(inversores_novos),
        
        # Totais
        'potencia_kwp': projeto.get('potencia_kwp', 0),
        'geracao_kwh_mes': projeto.get('geracao_kwh_mes', 0),
        'reducao_percentual': projeto.get('reducao_percentual', 0),
        'area_arranjos': projeto.get('area_arranjos', 0),
        'quantidade_modulos': projeto.get('quantidade_modulos', 0),
        
        # Data de geração
        'data_geracao': datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
    }
    
    return context

def generate_memorial_from_scratch(projeto):
    """Gera um memorial do zero se o template não existir"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Título
    title = doc.add_heading('MEMORIAL DESCRITIVO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Sistema Fotovoltaico Conectado à Rede', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Dados do cliente
    doc.add_heading('1. DADOS DO CLIENTE', level=1)
    
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'
    
    table.rows[0].cells[0].text = 'Nome do Cliente'
    table.rows[0].cells[1].text = projeto.get('nome_cliente', '')
    
    table.rows[1].cells[0].text = 'CPF/CNPJ'
    table.rows[1].cells[1].text = projeto.get('cpf_cnpj', '')
    
    table.rows[2].cells[0].text = 'Unidade Consumidora (UC)'
    table.rows[2].cells[1].text = projeto.get('uc', '')
    
    table.rows[3].cells[0].text = 'Endereço'
    table.rows[3].cells[1].text = projeto.get('endereco', '')
    
    table.rows[4].cells[0].text = 'Cidade/UF'
    table.rows[4].cells[1].text = f"{projeto.get('cidade', '')} / {projeto.get('uf', '')}"
    
    table.rows[5].cells[0].text = 'CEP'
    table.rows[5].cells[1].text = projeto.get('cep', '')
    
    table.rows[6].cells[0].text = 'Concessionária'
    table.rows[6].cells[1].text = projeto.get('concessionaria', '')
    
    table.rows[7].cells[0].text = 'Data do Projeto'
    table.rows[7].cells[1].text = projeto.get('data_projeto', '')
    
    # Tipo de projeto
    doc.add_heading('2. TIPO DE PROJETO', level=1)
    doc.add_paragraph(f"Tipo: {projeto.get('tipo_projeto', '')}")
    
    # Equipamentos
    doc.add_heading('3. EQUIPAMENTOS', level=1)
    
    if projeto.get('tipo_projeto') == 'Ampliação':
        doc.add_heading('Equipamentos Existentes', level=2)
        doc.add_paragraph(f"Módulos Existentes: {projeto.get('modulos_existentes', 0)}")
        doc.add_paragraph(f"Inversores Existentes: {projeto.get('inversores_existentes', '')}")
    
    if projeto.get('tipo_projeto') == 'Grid Zero':
        doc.add_heading('Equipamentos Grid Zero', level=2)
        doc.add_paragraph(f"Controlador: {projeto.get('controlador', '')}")
        doc.add_paragraph(f"Transdutor TC: {projeto.get('transdutor_tc', '')}")
        doc.add_paragraph(f"Chave Seccionadora: {projeto.get('chave_seccionadora', '')}")
    
    if projeto.get('tipo_projeto') == 'Art. 73-A':
        doc.add_heading('Dados Art. 73-A', level=2)
        doc.add_paragraph(f"Média de Consumo: {projeto.get('media_consumo', 0)} kWh")
        doc.add_paragraph(f"Fator de Carga: {projeto.get('fator_carga', 0)}")
        doc.add_paragraph(f"Fator de Ajuste: {projeto.get('fator_ajuste', 0)}")
    
    # Novo sistema
    doc.add_heading('Novo Sistema', level=2)
    doc.add_paragraph(f"Potência Total: {projeto.get('potencia_kwp', 0)} kWp")
    doc.add_paragraph(f"Geração Estimada: {projeto.get('geracao_kwh_mes', 0)} kWh/mês")
    doc.add_paragraph(f"Redução de Consumo: {projeto.get('reducao_percentual', 0)}%")
    doc.add_paragraph(f"Área de Arranjos: {projeto.get('area_arranjos', 0)} m²")
    doc.add_paragraph(f"Quantidade de Módulos: {projeto.get('quantidade_modulos', 0)}")
    
    # Rodapé
    doc.add_paragraph()
    footer = doc.add_paragraph(f"Documento gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Salvar
    output_folder = os.path.join(os.path.dirname(__file__), 'uploads')
    os.makedirs(output_folder, exist_ok=True)
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    nome_arquivo = f"Memorial_{projeto['nome_cliente']}_{timestamp}.docx"
    output_path = os.path.join(output_folder, nome_arquivo)
    
    doc.save(output_path)
    
    return output_path
